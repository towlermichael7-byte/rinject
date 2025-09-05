"""
Resume processing module for Resume Customizer application.
Coordinates all resume processing operations including parsing, document processing, and email handling.
"""

import concurrent.futures
from typing import List, Dict, Any, Optional, Callable
from io import BytesIO
from docx import Document
import streamlit as st

from text_parser import parse_input_text, get_parser
from document_processor import get_document_processor, FileProcessor
from email_handler import get_email_manager
from config import ERROR_MESSAGES, SUCCESS_MESSAGES, APP_CONFIG


class ResumeProcessor:
    """Main processor for individual resume operations."""
    
    def __init__(self):
        self.text_parser = get_parser()
        self.doc_processor = get_document_processor()
        self.file_processor = FileProcessor()
        
    def process_single_resume(
        self, 
        file_data: Dict[str, Any], 
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """
        Process a single resume with text parsing, document modification, and preparation.
        
        Args:
            file_data: Dictionary containing file information and settings
            progress_callback: Optional callback for progress updates
            
        Returns:
            Result dictionary with processing status and data
        """
        try:
            filename = file_data['filename']
            file_obj = file_data['file']
            raw_text = file_data['text']
            
            if progress_callback:
                progress_callback(f"Parsing tech stacks for {filename}...")
            
            # Parse tech stacks and points using the flexible parser
            selected_points, tech_stacks_used = self.text_parser.parse_tech_stacks(raw_text)
            
            if not selected_points or not tech_stacks_used:
                return {
                    'success': False, 
                    'error': ERROR_MESSAGES["no_tech_stacks"].format(filename=filename), 
                    'filename': filename
                }
            
            if progress_callback:
                progress_callback(f"Processing document for {filename}...")
            
            # Load and process document
            doc = Document(file_obj)
            projects_data = self.doc_processor.project_detector.find_projects_and_responsibilities(doc)
            
            if not projects_data:
                return {
                    'success': False, 
                    'error': ERROR_MESSAGES["no_projects"].format(filename=filename), 
                    'filename': filename
                }
            
            # Convert to structured format
            projects = []
            for i, project_tuple in enumerate(projects_data):
                if len(project_tuple) != 3:
                    logger.error(f"Malformed project tuple at index {i}: {project_tuple}")
                    continue
                title, start_idx, end_idx = project_tuple
                projects.append({
                    'title': title,
                    'index': i,
                    'responsibilities_start': start_idx,
                    'responsibilities_end': end_idx
                })
            
            # Distribute and add points using round-robin logic
            distribution_result = self.doc_processor.point_distributor.distribute_points_to_projects(
                projects, (selected_points, tech_stacks_used)
            )
            if not distribution_result['success']:
                return distribution_result
            # Add points to each project with mixed tech stacks
            total_added = 0
            # Sort projects by insertion point to process them in order
            sorted_projects = sorted(distribution_result['distribution'].items(),
                                   key=lambda x: x[1]['insertion_point'])
            
            # Keep track of how many paragraphs we've added to adjust insertion points
            paragraph_offset = 0
            
            for project_title, project_info in sorted_projects:
                # Adjust insertion point based on previous additions
                adjusted_project_info = project_info.copy()
                adjusted_project_info['insertion_point'] += paragraph_offset
                if 'responsibilities_end' in adjusted_project_info:
                    adjusted_project_info['responsibilities_end'] += paragraph_offset
                
                added = self.doc_processor.add_points_to_project(doc, adjusted_project_info)
                total_added += added
                
                # Update the offset for subsequent projects
                paragraph_offset += added
            points_added = total_added
            
            if progress_callback:
                progress_callback(f"Saving document for {filename}...")
            
            # Save to buffer
            output_buffer = BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            return {
                'success': True,
                'filename': filename,
                'buffer': output_buffer.getvalue(),
                'tech_stacks': tech_stacks_used,
                'points_added': points_added,
                'email_data': self._extract_email_data(file_data)
            }
            
        except Exception as e:
            return {
                'success': False, 
                'error': str(e), 
                'filename': file_data['filename']
            }
    
    def _extract_email_data(self, file_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Extract email configuration from file data.
        
        Args:
            file_data: File data dictionary
            
        Returns:
            Email configuration dictionary
        """
        return {
            'recipient': file_data.get('recipient_email', ''),
            'sender': file_data.get('sender_email', ''),
            'password': file_data.get('sender_password', ''),
            'smtp_server': file_data.get('smtp_server', ''),
            'smtp_port': file_data.get('smtp_port', 465),
            'subject': file_data.get('email_subject', ''),
            'body': file_data.get('email_body', '')
        }


class BulkResumeProcessor:
    """Handles bulk processing of multiple resumes with parallel execution."""
    
    def __init__(self):
        self.resume_processor = ResumeProcessor()
        self.file_processor = FileProcessor()
        self.email_manager = get_email_manager()
    
    def process_resumes_bulk(
        self, 
        files_data: List[Dict[str, Any]], 
        max_workers: int = None,
        progress_callback: Optional[Callable] = None
    ) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Process multiple resumes in parallel.
        
        Args:
            files_data: List of file data dictionaries
            max_workers: Maximum number of parallel workers
            progress_callback: Optional progress callback
            
        Returns:
            Tuple of (processed_resumes, failed_resumes)
        """
        processed_resumes = []
        failed_resumes = []
        
        if max_workers is None:
            max_workers = min(APP_CONFIG["max_workers_default"], len(files_data))
        
        # Optimization: Ensure all files have proper names before processing
        optimized_files_data = []
        for file_data in files_data:
            file_data['file'] = self.file_processor.ensure_file_has_name(
                file_data['file'], file_data['filename']
            )
            optimized_files_data.append(file_data)
        
        # Use ThreadPoolExecutor for parallel processing
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            # Submit all resume processing tasks
            future_to_file = {
                executor.submit(self.resume_processor.process_single_resume, file_data, progress_callback): 
                file_data['filename'] for file_data in optimized_files_data
            }
            
            # Collect results as they complete
            for future in concurrent.futures.as_completed(future_to_file):
                filename = future_to_file[future]
                try:
                    result = future.result()
                    if result['success']:
                        processed_resumes.append(result)
                    else:
                        failed_resumes.append(result)
                except Exception as e:
                    failed_resumes.append({
                        'success': False,
                        'filename': filename,
                        'error': str(e)
                    })
        
        # Clean up memory after bulk processing
        self.file_processor.cleanup_memory()
        
        return processed_resumes, failed_resumes
    
    def send_batch_emails(
        self, 
        processed_resumes: List[Dict[str, Any]], 
        progress_callback: Optional[Callable] = None
    ) -> List[Dict[str, Any]]:
        """
        Send emails for processed resumes.
        
        Args:
            processed_resumes: List of processed resume data
            progress_callback: Optional progress callback
            
        Returns:
            List of email sending results
        """
        return self.email_manager.send_batch_emails(processed_resumes, progress_callback)


class PreviewGenerator:
    """Generates preview of resume changes without modifying the original."""
    
    def __init__(self):
        self.text_parser = get_parser()
        self.doc_processor = get_document_processor()
    
    def generate_preview(
        self, 
        file_obj, 
        input_text: str, 
        manual_text: str = ""
    ) -> Dict[str, Any]:
        """
        Generate a preview of changes that will be made to the resume.
        
        Args:
            file_obj: Resume file object
            input_text: Tech stack input text
            manual_text: Optional manual points text
            
        Returns:
            Preview result dictionary
        """
        try:
            # Parse input text
            selected_points, tech_stacks_used = parse_input_text(input_text, manual_text)
            
            if not selected_points or not tech_stacks_used:
                return {
                    'success': False,
                    'error': 'Could not parse tech stacks from input'
                }
            
            # Load document and find projects
            doc = Document(file_obj)
            projects_data = self.doc_processor.project_detector.find_projects_and_responsibilities(doc)
            
            if not projects_data:
                return {
                    'success': False,
                    'error': 'No projects with Responsibilities sections found'
                }
            
            # Convert to structured format
            projects = []
            for i, (title, start_idx, end_idx) in enumerate(projects_data):
                projects.append({
                    'title': title,
                    'index': i,
                    'responsibilities_start': start_idx,
                    'responsibilities_end': end_idx
                })
            
            # Create preview document copy
            temp_buffer = BytesIO()
            doc.save(temp_buffer)
            temp_buffer.seek(0)
            
            preview_doc = Document(temp_buffer)
            preview_projects_data = self.doc_processor.project_detector.find_projects_and_responsibilities(preview_doc)
            
            # Convert to structured format with defensive check
            preview_projects = []
            for i, project_tuple in enumerate(preview_projects_data):
                if len(project_tuple) == 3:
                    title, start_idx, end_idx = project_tuple
                    preview_projects.append({
                        'title': title,
                        'index': i,
                        'responsibilities_start': start_idx,
                        'responsibilities_end': end_idx
                    })
                else:
                    logger.error(f"Malformed project tuple at index {i}: {project_tuple}")
                    continue
            
            # Apply changes to preview using round-robin logic
            distribution_result = self.doc_processor.point_distributor.distribute_points_to_projects(
                preview_projects, (selected_points, tech_stacks_used)
            )
            if not distribution_result['success']:
                return {'success': False, 'error': distribution_result['error']}
            
            # Add points to each project with mixed tech stacks for preview
            total_added = 0
            project_points_mapping = {}
            
            # Sort projects by insertion point to process them in order
            sorted_projects = sorted(distribution_result['distribution'].items(),
                                   key=lambda x: x[1]['insertion_point'])
            
            # Keep track of how many paragraphs we've added to adjust insertion points
            paragraph_offset = 0
            
            for project_title, project_info in sorted_projects:
                # Store which points are added to which project
                project_points = {}
                for tech_name, points in project_info['mixed_tech_stacks'].items():
                    project_points[tech_name] = points
                
                project_points_mapping[project_title] = project_points
                
                # Adjust insertion point based on previous additions
                adjusted_project_info = project_info.copy()
                adjusted_project_info['insertion_point'] += paragraph_offset
                if 'responsibilities_end' in adjusted_project_info:
                    adjusted_project_info['responsibilities_end'] += paragraph_offset
                
                # Add points to the document
                added = self.doc_processor.add_points_to_project(preview_doc, adjusted_project_info)
                total_added += added
                
                # Update the offset for subsequent projects
                paragraph_offset += added
            
            points_added = total_added
            
            # Generate preview content
            preview_content = self._extract_preview_content(preview_doc)
            
            return {
                'success': True,
                'points_added': points_added,
                'tech_stacks_used': tech_stacks_used,
                'selected_points': selected_points,
                'preview_content': preview_content,
                'preview_doc': preview_doc,
                'projects_count': len(projects),
                'project_points_mapping': project_points_mapping
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }
    
    def _extract_preview_content(self, doc: Document) -> str:
        """
        Extract readable content from document for preview.
        
        Args:
            doc: Document to extract content from
            
        Returns:
            Text content of the document
        """
        content_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                content_parts.append(text)
        
        return "\n\n".join(content_parts)


class ResumeManager:
    """Main manager class that coordinates all resume processing operations."""
    
    def __init__(self):
        self.resume_processor = ResumeProcessor()
        self.bulk_processor = BulkResumeProcessor()
        self.preview_generator = PreviewGenerator()
        self.email_manager = get_email_manager()
    
    def process_single_resume(
        self, 
        file_data: Dict[str, Any], 
        progress_callback: Optional[Callable] = None
    ) -> Dict[str, Any]:
        """Process a single resume."""
        return self.resume_processor.process_single_resume(file_data, progress_callback)
    
    def process_bulk_resumes(
        self, 
        files_data: List[Dict[str, Any]], 
        max_workers: int = None,
        progress_callback: Optional[Callable] = None
    ) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """Process multiple resumes in parallel."""
        return self.bulk_processor.process_resumes_bulk(files_data, max_workers, progress_callback)
    
    def generate_preview(
        self, 
        file_obj, 
        input_text: str, 
        manual_text: str = ""
    ) -> Dict[str, Any]:
        """Generate preview of resume changes."""
        return self.preview_generator.generate_preview(file_obj, input_text, manual_text)
    
    def send_batch_emails(
        self, 
        processed_resumes: List[Dict[str, Any]], 
        progress_callback: Optional[Callable] = None
    ) -> List[Dict[str, Any]]:
        """Send batch emails for processed resumes."""
        return self.bulk_processor.send_batch_emails(processed_resumes, progress_callback)
    
    def send_single_email(
        self, 
        smtp_server: str,
        smtp_port: int,
        sender_email: str,
        sender_password: str,
        recipient_email: str,
        subject: str,
        body: str,
        attachment_data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """Send a single email."""
        return self.email_manager.send_single_email(
            smtp_server, smtp_port, sender_email, sender_password,
            recipient_email, subject, body, attachment_data, filename
        )
    
    def validate_email_config(self, email_data: Dict[str, Any]) -> Dict[str, Any]:
        """Validate email configuration."""
        return self.email_manager.validate_email_config(email_data)
    
    def cleanup(self) -> None:
        """Clean up resources and connections."""
        self.email_manager.close_all_connections()


# Global resume manager instance
resume_manager = ResumeManager()


@st.cache_resource
def get_resume_manager() -> ResumeManager:
    """
    Get the global resume manager instance.
    
    Returns:
        ResumeManager instance
    """
    return resume_manager
