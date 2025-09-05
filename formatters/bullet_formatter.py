"""
Bullet point formatter for document processing.
Handles detection and formatting of bullet points in Word documents.
"""
from typing import Dict, Any, List, Optional, Tuple
from dataclasses import dataclass
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph

from .base_formatters import DocumentFormatter, ListFormatterMixin

@dataclass
class BulletFormatting:
    """Data class to hold bullet formatting information."""
    runs_formatting: List[Dict[str, Any]]
    paragraph_formatting: Dict[str, Any]
    style: str
    bullet_marker: str
    bullet_separator: str
    list_format: Dict[str, Any]


class BulletFormatter(DocumentFormatter, ListFormatterMixin):
    """Handles bullet point formatting and style preservation."""
    
    def __init__(self, bullet_markers: List[str] = None):
        """
        Initialize the BulletFormatter.
        
        Args:
            bullet_markers: List of bullet point markers to recognize
        """
        self.bullet_markers = bullet_markers or ['•', '●', '◦', '▪', '▫', '‣', '*', '-']
    
    def extract_formatting(self, doc: DocumentType, paragraph_index: int) -> Optional[BulletFormatting]:
        """
        Extract complete bullet formatting from a paragraph.
        
        Args:
            doc: Document containing the paragraph
            paragraph_index: Index of the paragraph to extract formatting from
            
        Returns:
            BulletFormatting object with formatting information, or None if not a bullet point
        """
        try:
            if paragraph_index >= len(doc.paragraphs):
                return None
                
            para = doc.paragraphs[paragraph_index]
            if not self._is_bullet_point(para.text):
                return None
                
            # Get basic formatting
            formatting_info = BulletFormatting(
                runs_formatting=[],
                paragraph_formatting={},
                style=para.style.name if para.style else 'Normal',
                bullet_marker=self._extract_bullet_marker(para.text) or '•',
                bullet_separator=self._detect_bullet_separator(para.text) or ' ',
                list_format=self._extract_list_format(para)
            )
            
            # Extract formatting from each run
            for run in para.runs:
                try:
                    run_format = {
                        'text': run.text,
                        'font_name': run.font.name if hasattr(run.font, 'name') else None,
                        'font_size': run.font.size.pt if hasattr(run.font, 'size') and run.font.size else None,
                        'bold': run.font.bold if hasattr(run.font, 'bold') else None,
                        'italic': run.font.italic if hasattr(run.font, 'italic') else None,
                        'underline': run.font.underline if hasattr(run.font, 'underline') else None,
                        'color': run.font.color.rgb if (hasattr(run.font, 'color') and 
                                                     run.font.color and 
                                                     hasattr(run.font.color, 'rgb')) else None
                    }
                    formatting_info.runs_formatting.append(run_format)
                except Exception:
                    continue
                    
            # Get paragraph formatting
            if hasattr(para, 'paragraph_format'):
                p_format = para.paragraph_format
                formatting_info.paragraph_formatting = {
                    'alignment': p_format.alignment,
                    'first_line_indent': p_format.first_line_indent,
                    'left_indent': p_format.left_indent,
                    'right_indent': p_format.right_indent,
                    'space_before': p_format.space_before,
                    'space_after': p_format.space_after,
                    'line_spacing': p_format.line_spacing,
                    'keep_together': p_format.keep_together,
                    'keep_with_next': p_format.keep_with_next,
                    'page_break_before': p_format.page_break_before,
                    'widow_control': p_format.widow_control
                }
                
            return formatting_info
            
        except Exception:
            # Return default formatting if anything fails
            return BulletFormatting(
                runs_formatting=[{'text': para.text if 'para' in locals() else ''}],
                paragraph_formatting={},
                style='Normal',
                bullet_marker='•',
                bullet_separator=' ',
                list_format={
                    'ilvl': 0,
                    'numId': 1,
                    'style': 'List Bullet',
                    'indent': 0,
                    'is_list': False
                }
            )
    
    def apply_formatting(self, paragraph: Paragraph, formatting: BulletFormatting, 
                        text: str, fallback_formatting: Optional[BulletFormatting] = None) -> None:
        """
        Apply extracted formatting to a new bullet point paragraph.
        
        Args:
            paragraph: Paragraph to format
            formatting: BulletFormatting object with formatting information
            text: Text content to add
            fallback_formatting: Optional fallback formatting if primary formatting fails
        """
        try:
            # Use fallback formatting if main formatting is missing or incomplete
            if not formatting and fallback_formatting:
                formatting = fallback_formatting

            # Apply paragraph style
            if formatting and formatting.style:
                try:
                    paragraph.style = formatting.style
                except Exception:
                    pass

            # Apply paragraph formatting
            if formatting and formatting.paragraph_formatting:
                pf = paragraph.paragraph_format
                for attr, value in formatting.paragraph_formatting.items():
                    if value is not None:
                        try:
                            setattr(pf, attr, value)
                        except Exception:
                            continue

            # Apply Word list formatting if available
            if formatting and formatting.list_format and formatting.list_format.get('is_list'):
                self._apply_list_formatting(paragraph, formatting.list_format)
            
            # Clear existing runs and add formatted text
            clean_text = self._clean_bullet_text(text)
            paragraph.clear()
            run = paragraph.add_run(clean_text)

            # Apply run formatting if available
            if formatting and formatting.runs_formatting:
                # Use the formatting from the first meaningful run
                primary_format = formatting.runs_formatting[0]
                for attr, value in primary_format.items():
                    if attr == 'text' or value is None:
                        continue
                    try:
                        if attr == 'color' and value:
                            run.font.color.rgb = value
                        else:
                            setattr(run.font, attr, value)
                    except Exception:
                        continue

        except Exception:
            # Fallback to basic formatting
            self._apply_basic_formatting(paragraph, text, formatting)
    
    def _extract_list_format(self, paragraph: Paragraph) -> Dict[str, Any]:
        """
        Extract list formatting information from a paragraph.
        
        Args:
            paragraph: The paragraph to extract list formatting from
            
        Returns:
            Dictionary containing list formatting information
        ""
        list_format = {
            'ilvl': 0,  # Default to top level
            'numId': 1,  # Default to first numbering ID
            'style': 'List Bullet',
            'indent': 0,
            'is_list': False
        }
        
        try:
            # Get style name safely
            if hasattr(paragraph, 'style') and paragraph.style is not None:
                list_format['style'] = paragraph.style.name
            
            # Get indentation
            if hasattr(paragraph, 'paragraph_format') and paragraph.paragraph_format is not None:
                if hasattr(paragraph.paragraph_format, 'left'):
                    list_format['indent'] = paragraph.paragraph_format.left or 0
            
            # Check if this is actually a bullet point
            if not self._is_bullet_point(paragraph.text):
                return list_format
                
            # If we get here, it's a bullet point
            list_format['is_list'] = True
            
            # Try to get Word list formatting if available
            if not hasattr(paragraph, '_element') or not hasattr(paragraph._element, 'pPr'):
                return list_format
                
            pPr = paragraph._element.pPr
            if pPr is None:
                return list_format
            
            numPr = getattr(pPr, 'numPr', None)
            if numPr is None:
                return list_format
            
            # Get list level (ilvl)
            if hasattr(numPr, 'ilvl') and hasattr(numPr.ilvl, 'val'):
                list_format['ilvl'] = numPr.ilvl.val
            
            # Get numbering ID (numId)
            if hasattr(numPr, 'numId') and hasattr(numPr.numId, 'val'):
                list_format['numId'] = numPr.numId.val
                        
        except Exception:
            pass
            
        return list_format
    
    def _is_bullet_point(self, text: str) -> bool:
        """Check if text starts with a bullet point marker."""
        text = text.strip()
        # Enhanced detection for various bullet formats
        return any(
            text.startswith(marker + ' ') or 
            text.startswith(marker) 
            for marker in self.bullet_markers
        ) or (text and text[0].isdigit() and '.' in text[:3])
    
    def _extract_bullet_marker(self, text: str) -> str:
        """Extract the bullet marker from text with enhanced detection."""
        text = text.strip()
        
        # Check priority markers first
        priority_markers = ['•', '●', '◦', '▪', '▫', '‣']
        for marker in priority_markers:
            if text.startswith(marker + '\t') or text.startswith(marker + ' ') or \
               (text.startswith(marker) and len(text) > 1):
                return marker
        
        # Check other markers
        for marker in self.bullet_markers:
            if text.startswith(marker + '\t') or text.startswith(marker + ' ') or \
               (text.startswith(marker) and len(text) > 1):
                return marker
                
        # Check for numbered bullets
        if text and text[0].isdigit():
            for i, char in enumerate(text):
                if char in '.)': 
                    return text[:i+1]
        
        return '•'  # Default to standard bullet
    
    def _detect_bullet_separator(self, text: str) -> str:
        """Detect whether bullet uses tab or space separator."""
        text = text.strip()
        for marker in self.bullet_markers:
            if text.startswith(marker + '\t'):
                return '\t'  # Tab separator
            elif text.startswith(marker + ' '):
                return ' '   # Space separator
        
        return '\t'  # Default to tab for better formatting
    
    def _clean_bullet_text(self, text: str) -> str:
        """Clean bullet text by removing any existing bullet markers."""
        # Remove common bullet markers from the beginning
        for marker in self.bullet_markers:
            if text.lstrip().startswith(marker):
                text = text.lstrip()[len(marker):].lstrip()
        
        # Also handle numbered bullets (e.g., "1. " or "1) ")
        if text and text[0].isdigit():
            for i, char in enumerate(text):
                if char in '.)':
                    text = text[i+1:].lstrip()
                    break
        
        return text
    
    def _apply_basic_formatting(self, paragraph: Paragraph, text: str, 
                              formatting: Optional[BulletFormatting] = None) -> None:
        """Apply basic bullet formatting as a fallback."""
        try:
            marker = formatting.bullet_marker if formatting and formatting.bullet_marker else '•'
            separator = formatting.bullet_separator if formatting and formatting.bullet_separator else '\t'
            clean_text = self._clean_bullet_text(text)
            paragraph.clear()
            paragraph.add_run(f"{marker}{separator}{clean_text}")
        except Exception:
            # Last resort: just add the text as-is
            paragraph.clear()
            paragraph.add_run(text)
