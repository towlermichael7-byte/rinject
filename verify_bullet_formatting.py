"""
Verify bullet formatting functionality.
"""
from docx import Document
from document_processor import BulletFormatter

def test_bullet_detection():
    """Test bullet point detection."""
    print("=== Testing Bullet Detection ===")
    formatter = BulletFormatter()
    
    test_cases = [
        ("• Test bullet", True, "•"),
        ("- Test bullet", True, "-"),
        ("* Test bullet", True, "*"),
        ("1. Test bullet", True, "1."),
        ("Just text", False, None),
        ("", False, None)
    ]
    
    for text, expected_result, expected_marker in test_cases:
        # Test bullet detection
        is_bullet = formatter._is_bullet_point(text)
        status = "PASS" if is_bullet == expected_result else "FAIL"
        print(f"{status}: {text!r} -> {'Bullet' if is_bullet else 'Not a bullet'}")
        
        # Test marker extraction
        if is_bullet:
            marker = formatter._extract_bullet_marker(text)
            status = "PASS" if marker == expected_marker else "FAIL"
            print(f"  {status}: Extracted marker: {marker!r}")

def test_bullet_formatting():
    """Test bullet formatting in a document."""
    print("\n=== Testing Bullet Formatting ===")
    
    # Create a new document
    doc = Document()
    doc.add_heading("Bullet Formatting Test", 0)
    
    # Add a test bullet point
    p = doc.add_paragraph("• Test bullet point")
    p.style = 'List Bullet'
    
    # Initialize formatter
    formatter = BulletFormatter()
    
    # Find the bullet point
    for i, para in enumerate(doc.paragraphs):
        if formatter._is_bullet_point(para.text):
            print(f"Found bullet point: {para.text}")
            
            # Get formatting
            formatting = formatter.get_bullet_formatting(doc, i)
            if formatting:
                print("Successfully got bullet formatting:")
                print(f"- Style: {formatting.get('style')}")
                print(f"- Marker: {formatting.get('bullet_marker')}")
                print(f"- Separator: {formatting.get('bullet_separator')!r}")
                
                # Try to apply formatting to a new paragraph
                new_para = doc.add_paragraph()
                formatter.apply_bullet_formatting(
                    new_para, 
                    formatting, 
                    "New bullet point with same formatting"
                )
                print("Added new bullet point with same formatting")
                break
    
    # Save the document
    output_path = "bullet_format_test.docx"
    doc.save(output_path)
    print(f"\nTest document saved to: {output_path}")

if __name__ == "__main__":
    test_bullet_detection()
    test_bullet_formatting()
    print("\nTest completed!")
