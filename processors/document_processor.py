"""
Main document processor for Resume Customizer.
Coordinates document processing operations including project detection and point distribution.
"""
import gc
import logging
from typing import Dict, List, Optional, Any, Tuple, Union, BinaryIO
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType

from formatters.bullet_formatter import BulletFormatter, BulletFormatting
from detectors.project_detector import ProjectDetector, ProjectInfo
from processors.point_distributor import PointDistributor, DistributionResult
from logger import get_logger

logger = get_logger(__name__)

class DocumentProcessor:
    """Main document processor that coordinates all operations."""
    
    def __init__(self, config: Optional[Dict] = None):
        """
        Initialize the DocumentProcessor.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or {}
        self.bullet_formatter = BulletFormatter(
            bullet_markers=self.config.get('bullet_markers', ['•', '●', '◦', '▪', '▫', '‣', '*', '-'])
        )
        self.project_detector = ProjectDetector(self.config.get('parsing', {}))
        self.point_distributor = PointDistributor()
    
    def process_document(
        self,
        file_data: Union[bytes, BinaryIO, str, Path],
        tech_stacks_data: Any,
        output_path: Optional[Union[str, Path]] = None
    ) -> Tuple[Optional[DocumentType], DistributionResult]:
        """
        Process a document by adding tech stack points to projects.
        
        Args:
            file_data: Input document as bytes, file-like object, or file path
            tech_stacks_data: Tech stack data to add to the document
            output_path: Optional path to save the processed document
            
        Returns:
            Tuple of (processed_document, distribution_result)
        """
        try:
            # Load the document
            doc = self._load_document(file_data)
            if not doc:
                raise ValueError("Failed to load document")
            
            # Detect projects in the document
            projects = self.project_detector.find_projects(doc)
            if not projects:
                logger.warning("No projects found in the document")
                return doc, DistributionResult()
            
            # Distribute points to projects
            distribution_result = self.point_distributor.distribute_points(
                [{"name": p.name, "start_index": p.start_index} for p in projects],
                tech_stacks_data
            )
            
            # Add points to the document
            self._add_points_to_document(doc, projects, distribution_result)
            
            # Save the document if output path is provided
            if output_path:
                self._save_document(doc, output_path)
            
            return doc, distribution_result
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}", exc_info=True)
            raise
        finally:
            # Clean up resources
            self.cleanup_resources()
    
    def _load_document(self, file_data: Union[bytes, BinaryIO, str, Path]) -> Optional[DocumentType]:
        """Load a document from various input types."""
        try:
            if isinstance(file_data, (str, Path)):
                return Document(file_data)
            elif isinstance(file_data, bytes):
                return Document(BytesIO(file_data))
            elif hasattr(file_data, 'read') and callable(file_data.read):
                # Reset file pointer to start if needed
                if hasattr(file_data, 'seek') and callable(file_data.seek):
                    file_data.seek(0)
                return Document(file_data)
            else:
                logger.error("Unsupported file data type")
                return None
        except Exception as e:
            logger.error(f"Error loading document: {str(e)}")
            return None
    
    def _save_document(self, doc: DocumentType, output_path: Union[str, Path]) -> bool:
        """Save the document to the specified path."""
        try:
            doc.save(output_path)
            return True
        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            return False
    
    def _add_points_to_document(
        self,
        doc: DocumentType,
        projects: List[ProjectInfo],
        distribution_result: DistributionResult
    ) -> None:
        """Add distributed points to the document."""
        if not distribution_result.distribution:
            return
        
        # Get the bullet formatting from the first bullet point in the document
        bullet_formatting = self._get_bullet_formatting(doc)
        
        # Add points to each project
        for project_name, points in distribution_result.distribution.items():
            if not points:
                continue
                
            # Find the project by name
            project = next((p for p in projects if p.name == project_name), None)
            if not project:
                continue
                
            # Add points after the project's end index
            self._add_points_to_project(
                doc=doc,
                project=project,
                points=points,
                bullet_formatting=bullet_formatting
            )
    
    def _get_bullet_formatting(self, doc: DocumentType) -> Optional[BulletFormatting]:
        """Get bullet formatting from the document."""
        # Look for the first bullet point in the document
        for i, para in enumerate(doc.paragraphs):
            if self.bullet_formatter._is_bullet_point(para.text):
                return self.bullet_formatter.extract_formatting(doc, i)
        return None
    
    def _add_points_to_project(
        self,
        doc: DocumentType,
        project: ProjectInfo,
        points: List[Dict[str, Any]],
        bullet_formatting: Optional[BulletFormatting] = None
    ) -> None:
        """Add points to a specific project in the document."""
        if not points:
            return
            
        # Find the insertion point (after the last bullet point of the project)
        insert_at = project.end_index + 1
        
        # Add each point
        for point in points:
            if not point or 'text' not in point:
                continue
                
            # Create a new paragraph for the bullet point
            new_para = doc.add_paragraph()
            
            # Apply bullet formatting
            if bullet_formatting:
                self.bullet_formatter.apply_formatting(
                    paragraph=new_para,
                    formatting=bullet_formatting,
                    text=point['text']
                )
            else:
                # Fallback to basic bullet point
                new_para.add_run(f"• {point['text']}")
            
            # Insert the new paragraph at the correct position
            doc.paragraphs.insert(insert_at, new_para._p)
            insert_at += 1
    
    def cleanup_resources(self) -> None:
        """Clean up resources and perform garbage collection."""
        gc.collect()


def get_document_processor(config: Optional[Dict] = None) -> DocumentProcessor:
    """
    Get a new instance of the document processor.
    
    Args:
        config: Optional configuration dictionary
        
    Returns:
        DocumentProcessor instance
    """
    return DocumentProcessor(config=config)


def cleanup_document_resources() -> None:
    """Clean up document processing resources."""
    gc.collect()


def force_memory_cleanup() -> None:
    """Force aggressive memory cleanup when usage is high."""
    import gc
    import os
    import psutil
    
    # Check memory usage
    process = psutil.Process(os.getpid())
    mem_info = process.memory_info()
    
    # If using more than 500MB, do aggressive cleanup
    if mem_info.rss > 500 * 1024 * 1024:  # 500MB
        logger.warning("High memory usage detected, forcing cleanup")
        gc.set_threshold(0)  # Disable garbage collection
        gc.collect()  # Run full collection
        gc.set_threshold(700, 10, 10)  # Restore default thresholds
        
        # Log memory after cleanup
        mem_info = process.memory_info()
        logger.info(f"Memory after cleanup: {mem_info.rss / (1024 * 1024):.2f}MB")
